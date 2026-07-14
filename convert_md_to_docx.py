import sys
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def parse_markdown_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set standard margins (1 inch on all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set normal style font to Times New Roman (standard for academic papers)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    # Read the markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    
    in_table = False
    table_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block, write it
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                in_code_block = False
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\n'))
            i += 1
            continue
            
        # Handle tables
        is_table_line = stripped.startswith('|') and stripped.endswith('|')
        if is_table_line:
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            # End of table, process and write it
            process_table(doc, table_lines)
            in_table = False
            table_lines = []
            # Fall through to process current line if not empty
            if not stripped:
                i += 1
                continue
                
        # Handle Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            i += 1
            continue
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
            i += 1
            continue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)
            i += 1
            continue
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[5:])
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue
        elif stripped.startswith('##### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[6:])
            run.italic = True
            run.font.size = Pt(12)
            i += 1
            continue

        # Handle horizontal rules
        if stripped == '---':
            doc.add_page_break()
            i += 1
            continue
            
        # Handle Block Equations starting with $$
        if stripped.startswith('$$') and stripped.endswith('$$'):
            eq_text = stripped.strip('$').strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(eq_text)
            run.italic = True
            i += 1
            continue

        # Handle lists (bullet/number)
        is_bullet = stripped.startswith('* ') or stripped.startswith('- ')
        is_numbered = re.match(r'^\d+\.\s', stripped)
        
        if is_bullet:
            content = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, content)
            i += 1
            continue
        elif is_numbered:
            match = is_numbered.group(0)
            content = stripped[len(match):]
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            parse_inline_formatting(p, content)
            i += 1
            continue
            
        # Handle ordinary paragraphs
        if stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, stripped)
            
        i += 1

    doc.save(docx_path)
    print(f"[SUCCESS] Saved docx to: {docx_path}")

def parse_inline_formatting(paragraph, text):
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$\$.*?\$\$(?!.*?\$))', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        elif token.startswith('$$') and token.endswith('$$'):
            run = paragraph.add_run(token[2:-2])
            run.italic = True
        else:
            paragraph.add_run(token)

def process_table(doc, table_lines):
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if all(re.match(r'^:?-+:?$', c) for c in cells):
            continue
        rows_data.append(cells)
        
    if not rows_data:
        return
        
    num_cols = len(rows_data[0])
    num_rows = len(rows_data)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            
            run = p.add_run(val)
            run.font.size = Pt(10)
            if r_idx == 0:
                run.bold = True
                
    doc.add_paragraph().paragraph_format.space_before = Pt(6)

def convert_to_txt(md_path, txt_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"[SUCCESS] Saved txt to: {txt_path}")

if __name__ == '__main__':
    md = r"D:\THESIS ZAKY\Jurnal Output 2\Ethereum_Smart_Contract_Cocoa_Traceability.md"
    docx = r"D:\THESIS ZAKY\Jurnal Output 2\Ethereum_Smart_Contract_Cocoa_Traceability.docx"
    txt = r"D:\THESIS ZAKY\Jurnal Output 2\Ethereum_Smart_Contract_Cocoa_Traceability.txt"
    
    parse_markdown_to_docx(md, docx)
    convert_to_txt(md, txt)
